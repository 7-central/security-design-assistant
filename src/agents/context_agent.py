"""Context Agent for processing context documents using Google GenAI SDK."""
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from src.agents.base_agent_v2 import BaseAgentV2

logger = logging.getLogger(__name__)


class ContextAgent(BaseAgentV2):
    """Agent for processing context documents (specifications, requirements, etc.)."""

    def __init__(self, storage, job):
        """Initialize Context Agent with Gemini 2.5 Pro."""
        super().__init__(storage, job)
        self.model_name = "models/gemini-2.5-pro"  # Using consistent model across agents

    async def extract_context(self, file_path: str | Path) -> dict[str, Any]:
        """Process DOCX/PDF/text file into structured JSON format.

        Args:
            file_path: Path to the context file

        Returns:
            Dictionary with structured context data
        """
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()

        if file_extension == '.docx':
            return await self._process_docx_file(file_path)
        elif file_extension == '.pdf':
            return await self._process_pdf_file(file_path)
        elif file_extension in ['.txt', '.text']:
            return await self._process_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    async def summarize_specifications(self, raw_text: str) -> dict[str, Any]:
        """Extract lock types and requirements from raw text.

        Args:
            raw_text: Raw text content to summarize

        Returns:
            Dictionary with extracted specifications
        """
        prompt = self._build_summarization_prompt(raw_text)

        try:
            response = await self._generate_with_retry(prompt, model_name=self.model_name)
            content = response.text if hasattr(response, 'text') else str(response)
            return self._parse_json_response(content)
        except Exception as e:
            logger.error(f"Error summarizing specifications: {e}")
            return {"sections": [], "error": str(e)}

    async def _process_docx_file(self, file_path: Path) -> dict[str, Any]:
        """Process DOCX file using python-docx library.

        Args:
            file_path: Path to DOCX file

        Returns:
            Structured context data
        """
        try:
            try:
                import docx
            except ImportError as e:
                logger.error("python-docx library not installed")
                raise ImportError("python-docx library is required for DOCX processing") from e

            doc = docx.Document(str(file_path))

            # Extract text with section headings
            sections = []
            current_section = {"title": "Document Content", "content": [], "type": "general"}

            for paragraph in doc.paragraphs:
                # Check if it's a heading
                if paragraph.style.name.startswith('Heading'):
                    # Save current section if it has content
                    if current_section["content"]:
                        current_section["content"] = "\n".join(current_section["content"])
                        sections.append(current_section)

                    # Start new section
                    section_type = "specification" if any(
                        keyword in paragraph.text.lower()
                        for keyword in ["specification", "requirement", "lock", "hardware", "standard"]
                    ) else "general"

                    current_section = {
                        "title": paragraph.text,
                        "content": [],
                        "type": section_type
                    }
                else:
                    # Add to current section
                    if paragraph.text.strip():
                        current_section["content"].append(paragraph.text)

            # Add last section
            if current_section["content"]:
                current_section["content"] = "\n".join(current_section["content"])
                sections.append(current_section)

            # Extract tables
            for i, table in enumerate(doc.tables):
                table_content = self._extract_table_content(table)
                if table_content:
                    sections.append({
                        "title": f"Table {i + 1}",
                        "content": table_content,
                        "type": "specification"
                    })

            # Use Gemini Flash to further structure the content
            full_text = "\n\n".join([f"{s['title']}:\n{s['content']}" for s in sections])
            structured_data = await self.summarize_specifications(full_text)

            # Merge extracted sections with AI-structured data
            if "sections" in structured_data:
                return structured_data
            else:
                return {"sections": sections}
        except Exception as e:
            logger.error(f"Error processing DOCX file: {e}")
            return {"sections": [], "error": str(e)}

    async def _process_pdf_file(self, file_path: Path) -> dict[str, Any]:
        """Process PDF file using Gemini Flash multimodal.

        Args:
            file_path: Path to PDF file

        Returns:
            Structured context data
        """
        try:
            # Upload PDF file to Gemini
            uploaded_file = self.upload_file(str(file_path))

            # Build multimodal prompt
            prompt = """Analyze this PDF document and extract structured information for a security access control system.

Extract the following:
1. Lock types and specifications (especially types 11-22 if present)
2. Hardware requirements and standards
3. General project requirements
4. Special instructions or exceptions

Return a JSON object with this structure:
{
    "sections": [
        {
            "title": "Section Title",
            "content": "Extracted content",
            "type": "specification" or "general"
        }
    ]
}"""

            # Generate content with uploaded file
            response = self.generate_content(
                model_name=self.model_name,
                contents=[prompt, uploaded_file],
                generation_config={
                    "temperature": 0.1,
                    "top_p": 0.95,
                    "max_output_tokens": 8192,
                }
            )

            content = response.text if hasattr(response, 'text') else str(response)
            return self._parse_json_response(content)

        except Exception as e:
            logger.error(f"Error processing PDF file: {e}")
            return {"sections": [], "error": str(e)}

    async def _process_text_file(self, file_path: Path) -> dict[str, Any]:
        """Process text file.

        Args:
            file_path: Path to text file

        Returns:
            Structured context data
        """
        try:
            with open(file_path, encoding='utf-8') as f:
                content = f.read()

            return await self.summarize_specifications(content)

        except Exception as e:
            logger.error(f"Error processing text file: {e}")
            return {"sections": [], "error": str(e)}

    def _extract_table_content(self, table) -> str:
        """Extract content from a DOCX table.

        Args:
            table: DOCX table object

        Returns:
            Formatted table content as string
        """
        rows = []
        for row in table.rows:
            row_content = [cell.text.strip() for cell in row.cells]
            if any(row_content):  # Skip empty rows
                rows.append(" | ".join(row_content))

        return "\n".join(rows) if rows else ""

    def _build_summarization_prompt(self, raw_text: str) -> str:
        """Build prompt for text summarization.

        Args:
            raw_text: Raw text to summarize

        Returns:
            Formatted prompt string
        """
        return f"""You are analyzing a context document for a security drawing processing system.
Your task is to extract structured information that will help understand access control components.

Extract the following information:
1. Lock types and specifications (pay special attention to types 11-22)
2. Hardware requirements and standards
3. General project information
4. Special instructions or exceptions

Document Content:
{raw_text[:10000]}  # Limit to first 10k chars to avoid token limits

Return a JSON object with this structure:
{{
    "sections": [
        {{
            "title": "Section Title",
            "content": "Extracted content",
            "type": "specification" or "general"
        }}
    ]
}}"""

    def _parse_json_response(self, response: str) -> dict[str, Any]:
        """Parse JSON response from AI model.

        Args:
            response: Raw response string

        Returns:
            Parsed JSON dictionary
        """
        try:
            # Extract JSON from response
            if "```json" in response:
                json_str = response.split("```json")[1].split("```")[0].strip()
            elif "{" in response and "}" in response:
                start = response.find("{")
                end = response.rfind("}") + 1
                json_str = response[start:end]
            else:
                json_str = response

            return json.loads(json_str)

        except json.JSONDecodeError as e:
            logger.warning(f"Failed to parse JSON response: {e}")
            # Return a basic structure with raw content
            return {
                "sections": [{
                    "title": "Raw Content",
                    "content": response[:5000],  # Limit size
                    "type": "general"
                }]
            }

    async def process(self, input_data: dict[str, Any]) -> dict[str, Any]:
        """Process context document and extract relevant information.

        Args:
            input_data: Dictionary containing:
                - context_file_path: Path to context file (optional)
                - context_text: Raw context text (optional)
                - context_type: Type of context (from classifier)

        Returns:
            Dictionary with extracted context information matching schema
        """
        start_time = datetime.utcnow()

        self.log_structured("info", "Starting context processing",
                          context_type=input_data.get("context_type"))

        try:
            # Check if context is provided
            if not input_data.get("context_file_path") and not input_data.get("context_text"):
                self.log_structured("info", "No context provided, skipping context processing")
                return {
                    "sections": [],
                    "metadata": {
                        "source_type": "none",
                        "sections_count": 0,
                        "tokens_used": 0,
                        "processing_time_ms": 0
                    }
                }

            # Process based on input type
            context_data = {"sections": []}
            tokens_used = 0

            if input_data.get("context_file_path"):
                file_path = input_data["context_file_path"]
                context_data = await self.extract_context(file_path)

                # Estimate tokens based on file size
                file_size = Path(file_path).stat().st_size
                tokens_used = file_size // 4  # Rough estimate

            elif input_data.get("context_text"):
                context_data = await self.summarize_specifications(input_data["context_text"])
                tokens_used = self.estimate_tokens(input_data["context_text"])

            # Calculate processing time
            processing_time = (datetime.utcnow() - start_time).total_seconds() * 1000

            # Build result matching the required schema
            result = {
                "sections": context_data.get("sections", []),
                "metadata": {
                    "source_type": input_data.get("context_type", {}).get("type", "unknown"),
                    "sections_count": len(context_data.get("sections", [])),
                    "tokens_used": tokens_used,
                    "processing_time_ms": int(processing_time)
                }
            }

            # Log metrics
            self.log_structured("info", "Context processing complete",
                              sections_found=result["metadata"]["sections_count"],
                              tokens_used=tokens_used,
                              processing_time_ms=int(processing_time))

            # Save checkpoint
            await self.save_checkpoint("context", result)

            # Update job metadata
            self.job.update_metadata({
                "context_processing": {
                    "type": result["metadata"]["source_type"],
                    "sections_found": result["metadata"]["sections_count"],
                    "tokens_used": tokens_used
                }
            })

            return result

        except Exception as e:
            processing_time = (datetime.utcnow() - start_time).total_seconds() * 1000

            self.log_structured("error", f"Context processing failed: {e!s}",
                              error_type=type(e).__name__)

            # Return empty context on failure (graceful degradation)
            return {
                "sections": [],
                "metadata": {
                    "source_type": input_data.get("context_type", {}).get("type", "unknown"),
                    "sections_count": 0,
                    "tokens_used": 0,
                    "processing_time_ms": int(processing_time),
                    "error": str(e)
                }
            }
